'''
[python处理word] - docx
    1. 后缀需为.docx
    2. 需要将原文设置好样式，比如某一段是<正文>，则需要设置<正文>这一整个template的样式为期望的样式

    ref：
        https://www.mscto.com/python/606892.html
'''
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

document = Document(r"doc/template.docx")
paragraphs = document.paragraphs

name, studentID, academy = "张喆", "1754060", "软件学院"

'''
通过python的正则表达式，可以很简单地实现文本的替换和查找
    将原文中的academy占位符替换成动态数据
    （但是这个操作会丢失掉行内信息，加粗、下划线等等）
'''
text = re.sub('academy', academy, paragraphs[4].text)
paragraphs[4].text = text


'''
重新写所有的文本信息，可以制造行内样式
'''
paragraphs[4].text = ""
paragraphs[4].add_run("兹证明:同济大学 ")
paragraphs[4].add_run(" " + academy + " ").underline = True
paragraphs[4].add_run(" 学院学生：姓名")
paragraphs[4].add_run(" " + name + " ").underline = True
paragraphs[4].add_run("，学号 ")
paragraphs[4].add_run(" " + studentID + " ").underline = True
paragraphs[4].add_run("，参与同济大学2020年大学生暑期社会实践活动。")


document.save("doc/refactor.docx")